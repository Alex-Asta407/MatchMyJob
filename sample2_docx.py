# generate_test_docx.py

from docx import Document

doc = Document()
doc.add_heading("Jane Smith", level=1)

doc.add_paragraph("SKILLS:")
doc.add_paragraph("Java; Spring Boot; Hibernate")
doc.add_paragraph("JavaScript; React.js; Node.js")

doc.add_paragraph("EXPERIENCE:")
doc.add_paragraph("2018–Present  Senior Backend Developer, TechSolutions")
doc.add_paragraph("• Designed microservices with Spring Boot")
doc.add_paragraph("• Integrated OAuth2 authentication")
doc.add_paragraph("• Reduced response time by 40%")

doc.add_paragraph("2016–2018  Junior Developer, WebWorks")
doc.add_paragraph("• Built SPAs in React.js")
doc.add_paragraph("• Wrote Node.js serverless functions")

doc.save("sample2.docx")
print("Wrote sample2.docx")
